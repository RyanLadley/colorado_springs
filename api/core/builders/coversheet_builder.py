#!env/bin/python3
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.shared import Cm

from datetime import datetime
from io import BytesIO

import api.core.utilities as utilities

def build_single_invoice(coversheet):

    document = Document()
    section = document.sections[0]
    section.left_margin, section.right_margin = (Inches(.75), Inches(.75))

    title = document.add_paragraph('PPRTA Invoice Cover Sheet\rPublic Works- Operations and Maintenance Division')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bold_style = []
    underline_style = []

    invoice_paragraph = document.add_paragraph()
    bold_style.append(invoice_paragraph.add_run("Invoice No.:\t"))
    underline_style.append(invoice_paragraph.add_run("\t" +coversheet.invoice_no +"\t"))
    bold_style.append(invoice_paragraph.add_run("\tfor\t"))
    underline_style.append(invoice_paragraph.add_run("\t" +coversheet.vendor.name +"\t"))

    pprta_paragraph = document.add_paragraph()
    bold_style.append(pprta_paragraph.add_run("PPRTA Contract/PO No.:\t"))
    underline_style.append(pprta_paragraph.add_run("\t" +coversheet.vendor.contract_no +"\t"))
 
    description_paragraph = document.add_paragraph()
    bold_style.append(description_paragraph.add_run("Description:\t"))
    underline_style.append(description_paragraph.add_run("\t" +coversheet.description +"\t"))


    _add_account_table(document, coversheet.transactions)
    _add_approvals(document)
    _add_notes(document)


    for style in bold_style:
        style.bold = True

    for style in underline_style:
        style.underline = True

    document.add_page_break()

    #file = BytesIO()
    file_name = datetime.now().strftime('%Y-%m-%d_%H-%M') +".docx"
    document.save('api/exports/coversheets/single_invoice/' +file_name)

    return file_name


def build_project_sheet(coversheet):

    document = Document()
    section = document.sections[0]
    section.left_margin, section.right_margin = (Inches(.75), Inches(.75))

    title = document.add_paragraph('PPRTA Invoice Cover Sheet\rPublic Works- Operations and Maintenance Division')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_project_table(document, coversheet.transactions)
    document.add_paragraph()
    _add_contract_table(document, coversheet.vendor, coversheet.description)
    document.add_paragraph()
    _add_account_table(document, coversheet.transactions)
    _add_approvals(document)
    _add_notes(document)

    document.add_page_break()

    #file = BytesIO()
    file_name = datetime.now().strftime('%Y-%m-%d_%H-%M') +".docx"
    document.save('api/exports/coversheets/project/' +file_name)

    return file_name

def _add_account_table(document, transactions):
    
    table = document.add_table(rows = 1, cols = 4)
    table.style = None

    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(6)
    table.columns[2].width = Cm(4)
    table.columns[3].width = Cm(3)

    table.autofit = True
    header = table.rows[0].cells
    header[0].paragraphs[0].add_run("Project or Program").bold = True
    header[1].paragraphs[0].add_run("City Account Code").bold = True
    header[2].paragraphs[0].add_run("PPRTA Account Code").bold = True
    header[3].paragraphs[0].add_run("Amount").bold = True

    for cell in header:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    total = 0
    for transaction in transactions:
        for account in transaction.city_account_assignments:
            row = table.add_row().cells
            row[0].text = transaction.pprta_codes.project_description
            row[1].text = "{}-{}-{}-{}".format(
                account.city_account_no, transaction.pprta_codes.fund_no, transaction.pprta_codes.dept_no, transaction.pprta_codes.project_no)
            row[2].text = "{}-{}".format(
                transaction.pprta_codes.account_prefix, transaction.pprta_codes.account_no)
            row[3].text = utilities.currency(account.amount)
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            total += account.amount


    total_row = table.add_row().cells
    total_row[2].paragraphs[0].add_run("Total").bold = True
    total_row[3].paragraphs[0].add_run(utilities.currency(total)).bold = True   
    total_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_approvals(document):

    approvals = ["Procurements Svcs Approval", "Street Division Approval", "Finance Dept Approval\t"]
    for approval in approvals:
        document.add_paragraph()
        paragraph = document.add_paragraph()
        paragraph.add_run(approval +'\t')
        paragraph.add_run("\t\t\t\t\t").underline = True
        paragraph.add_run("\tDate: ")
        paragraph.add_run("\t\t\t").underline = True


def _add_notes(document):
    
    table = document.add_table(rows = 4, cols = 1)
    table.style = 'Table Grid'

    paragraph = table.rows[0].cells[0].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notes = paragraph.add_run("Notes")
    notes.bold = True
    notes.underline = True


def _add_vendor_table(document):

    table = document.add_table(rows = 3, cols = 2)

    invoice_cell = table.rows[0].cells[0]
    company_cell = table.rows[0].cells[1]
    pprta_cell = table.rows[1].cells[0]
    contract_cell = table.rows[1].cells[1]
    description_cell = table.rows[2].cells[0]

    description_cell.merge(table.rows[2].cells[1])

    invoice_cell.text = "Invoice No."
    invoice_cell.add_paragraph("910018350, 9100185275")

    company_cell.text = "Company Name"
    company_cell.add_paragraph("Kiewit")

    pprta_cell.text = "PPRTA No."
    pprta_cell.add_paragraph("")

    contract_cell.text = "Contract No."
    contract_cell.add_paragraph("T006731")

    description_cell.text = "Description"
    description_cell.add_paragraph("Asphalt and track for road repair")


def _add_project_table(document, transactions):

    table = document.add_table(rows = 1, cols = 4)
    table.style = 'TableGrid'

    title_cells = table.rows[0].cells
    title_start = title_cells[0]
    title_end = title_cells[3]

    title_start.merge(title_end)
    title_start.paragraphs[0].add_run(transactions[0].vendor_name).bold = True
    title_start.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table.add_row()
    header = table.rows[1].cells
    header[0].paragraphs[0].add_run("Date").bold = True
    header[1].paragraphs[0].add_run("Invoice #").bold = True
    header[2].paragraphs[0].add_run(transactions[0].pprta_codes.project_description).bold = True
    header[3].paragraphs[0].add_run("Total").bold = True

    for cell in header:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    project_total = 0
    transaction_total = 0
    for transaction in transactions:
        row = table.add_row().cells
        row[0].text = transaction.invoice_date
        row[1].text = transaction.invoice_no

        row[2].text = utilities.currency(transaction.expense)
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row[3].text = utilities.currency(transaction.expense)
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        project_total += transaction.expense
        transaction_total += transaction.expense

    total_row = table.add_row().cells
    total_label = total_row[0]
    total_label.merge(total_row[1])

    total_label.paragraphs[0].add_run("Total Amount").bold = True
    
    total_row[2].text = utilities.currency(project_total)
    total_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    total_row[3].text = utilities.currency(transaction_total)
    total_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_contract_table(document, vendor, description):

    table = document.add_table(rows = 2, cols = 1)

    contract_row = table.rows[0].cells
    contract_row[0].paragraphs[0].add_run("Contract Number: ").bold = True
    contract_row[0].paragraphs[0].add_run(vendor.contract_no)
    contract_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    description_row = table.rows[1].cells
    description_row[0].text = " "
    description_row[0].paragraphs[0].add_run("Description: ").bold = True
    description_row[0].paragraphs[0].add_run(description)

